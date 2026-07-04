from __future__ import annotations

from pathlib import Path
from docx import Document


def sanitize_relationship_type(raw_type: str) -> str:
    """
    Cypher relationship types can't be passed as query parameters, so they
    have to be interpolated into the query string directly. Since raw_type
    usually comes from LLM-extracted text, sanitize it to a safe
    UPPER_SNAKE_CASE identifier before interpolation, to avoid Cypher
    injection or syntax errors from unexpected characters.
    """
    safe = "".join(ch if ch.isalnum() else "_" for ch in raw_type.strip())
    safe = safe.upper().strip("_") or "RELATIONSHIP"
    if safe[0].isdigit():
        safe = f"REL_{safe}"
    return safe



def load_docx_text(file_path: str | Path) -> str:
    """
    Load a .docx file and return its paragraph text joined by newlines.

    This mirrors the original inline logic:
        doc = Document(file_path)
        raw_data = "\\n".join(p.text for p in doc.paragraphs)

    Empty paragraphs are preserved (not filtered) to keep paragraph
    breaks intact, which matters for downstream extraction chunking
    (e.g. langextract's max_char_buffer splitting).

    Raises:
        FileNotFoundError: if `file_path` does not exist.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    doc = Document(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()